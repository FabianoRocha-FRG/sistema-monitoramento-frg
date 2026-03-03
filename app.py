import streamlit as st
import requests
import pandas as pd
import sqlite3
from docx import Document
from datetime import datetime

# -----------------------------
# CONFIGURAÇÃO DA PÁGINA
# -----------------------------
st.set_page_config(
    page_title="Atlas Legislativo FRG",
    page_icon="🗺️",
    layout="wide"
)

st.title("🗺️ Atlas Legislativo FRG")
st.markdown("### Plataforma de Inteligência e Monitoramento Legislativo Federal")
st.divider()

# -----------------------------
# BANCO DE DADOS LOCAL
# -----------------------------
conn = sqlite3.connect("frg_monitoramento.db", check_same_thread=False)
cursor = conn.cursor()

cursor.execute("""
CREATE TABLE IF NOT EXISTS radar (
    id TEXT,
    casa TEXT,
    tipo TEXT,
    numero TEXT,
    ano TEXT,
    ementa TEXT,
    risco TEXT,
    cliente TEXT,
    status TEXT,
    prioridade TEXT,
    responsavel TEXT,
    proxima_acao TEXT,
    data_salvo TEXT
)
""")

conn.commit()

# -----------------------------
# CLASSIFICAÇÃO DE RISCO
# -----------------------------
def classificar_risco(ementa):
    palavras_criticas = ["tribut", "imposto", "regulação", "taxa", "penal"]
    palavras_altas = ["obrigação", "proibição", "multa"]

    ementa_lower = ementa.lower()

    if any(p in ementa_lower for p in palavras_criticas):
        return "🔴 Crítico"
    elif any(p in ementa_lower for p in palavras_altas):
        return "🟠 Alto"
    else:
        return "🟡 Médio"

# -----------------------------
# FUNÇÕES DE BUSCA
# -----------------------------
def buscar_camara(palavra):
    url = "https://dadosabertos.camara.leg.br/api/v2/proposicoes?keywords=" + palavra
    response = requests.get(url)

    if response.status_code == 200:
        dados = response.json().get("dados", [])
        if dados:
            df = pd.DataFrame(dados)
            df = df[["id","siglaTipo","numero","ano","ementa"]]
            df["casa"] = "Câmara"
            return df
    return pd.DataFrame()

def buscar_senado(palavra):
    url = f"https://legis.senado.leg.br/dadosabertos/materia/pesquisa/lista.json?q={palavra}"
    response = requests.get(url)

    if response.status_code == 200:
        try:
            materias = response.json()["PesquisaBasicaMateria"]["Materias"]["Materia"]
            df = pd.json_normalize(materias)
            df = df[["CodigoMateria","SiglaSubtipoMateria","NumeroMateria","AnoMateria","EmentaMateria"]]
            df.columns = ["id","tipo","numero","ano","ementa"]
            df["casa"] = "Senado"
            return df
        except:
            return pd.DataFrame()
    return pd.DataFrame()

# -----------------------------
# INTERFACE
# -----------------------------
modo = st.selectbox("Modo de Operação", [
    "Busca Unificada",
    "Câmara",
    "Senado",
    "Radar Estratégico FRG"
])

if modo != "Radar Estratégico FRG":

    palavra = st.text_input("🔎 Palavra-chave para monitoramento")

    if st.button("Buscar"):

        if palavra:

            df_camara = buscar_camara(palavra)
            df_senado = buscar_senado(palavra)

            if modo == "Câmara":
                df = df_camara
            elif modo == "Senado":
                df = df_senado
            else:
                df = pd.concat([df_camara, df_senado])

            if not df.empty:

                df["risco"] = df["ementa"].apply(classificar_risco)

                st.dataframe(df, use_container_width=True)

                st.subheader("Salvar no Radar Estratégico")

                cliente = st.text_input("Cliente vinculado")
                status = st.selectbox("Status", ["Monitorando", "Em articulação", "Reunião realizada", "Nota técnica enviada"])
                prioridade = st.selectbox("Prioridade", ["Baixa", "Média", "Alta", "Urgente"])
                responsavel = st.text_input("Responsável FRG")
                proxima_acao = st.text_input("Próxima ação prevista")

                if st.button("Salvar no Radar"):
                    for _, row in df.iterrows():
                        cursor.execute("""
                        INSERT INTO radar VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """, (
                            str(row["id"]),
                            row["casa"],
                            row.get("siglaTipo", row.get("tipo","")),
                            str(row["numero"]),
                            str(row["ano"]),
                            row["ementa"],
                            row["risco"],
                            cliente,
                            status,
                            prioridade,
                            responsavel,
                            proxima_acao,
                            datetime.now().strftime("%d/%m/%Y")
                        ))
                    conn.commit()
                    st.success("Projeto salvo no Radar Estratégico FRG")

            else:
                st.warning("Nenhuma proposição encontrada.")

# -----------------------------
# RADAR
# -----------------------------
if modo == "Radar Estratégico FRG":

    df_radar = pd.read_sql_query("SELECT * FROM radar", conn)

    if not df_radar.empty:

        st.subheader("📊 Dashboard Executivo")

        total = len(df_radar)
        criticos = len(df_radar[df_radar["risco"] == "🔴 Crítico"])
        altos = len(df_radar[df_radar["risco"] == "🟠 Alto"])

        col1, col2, col3 = st.columns(3)
        col1.metric("Total Monitorados", total)
        col2.metric("Projetos Críticos", criticos)
        col3.metric("Projetos Alto Risco", altos)

        st.divider()
        st.dataframe(df_radar, use_container_width=True)

        cliente_relatorio = st.text_input("Gerar relatório para qual cliente?")

        if st.button("Gerar Relatório Word"):
            doc = Document()
            doc.add_heading("Atlas Legislativo FRG", level=1)
            doc.add_paragraph(f"Cliente: {cliente_relatorio}")
            doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}")

            for _, row in df_radar.iterrows():
                if row["cliente"] == cliente_relatorio:
                    doc.add_heading(f"{row['tipo']} {row['numero']}/{row['ano']}", level=2)
                    doc.add_paragraph(row["ementa"])
                    doc.add_paragraph(f"Risco: {row['risco']}")
                    doc.add_paragraph("")

            nome_arquivo = f"Atlas_Legislativo_FRG_{cliente_relatorio}.docx"
            doc.save(nome_arquivo)

            st.success(f"Relatório gerado: {nome_arquivo}")

    else:
        st.info("Radar Estratégico vazio.")